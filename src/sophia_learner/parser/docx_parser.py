"""
DOCX Parser Module

Provides parser for Microsoft Word .docx files using python-docx library
for text extraction and metadata retrieval.
"""

from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .base_parser import BaseParser, ParserError, EncryptedFileError
from ..security.sandbox import Sandbox


class DocxParser(BaseParser):
    """
    Parser for Microsoft Word .docx files.
    
    Uses python-docx library to extract text content including paragraphs,
    tables, headers, and footers. Also extracts document metadata from
    core properties.
    
    Attributes:
        extract_headers: Whether to extract header content (default: True)
        extract_footers: Whether to extract footer content (default: True)
        extract_tables: Whether to extract table content (default: True)
        sandbox: Optional sandbox for resource-limited execution
    """
    
    def __init__(
        self, 
        sandbox: Optional[Sandbox] = None,
        extract_headers: bool = True,
        extract_footers: bool = True,
        extract_tables: bool = True
    ):
        """
        Initialize the DOCX parser.
        
        Args:
            sandbox: Optional Sandbox instance for resource-limited execution
            extract_headers: Whether to include header content in extraction
            extract_footers: Whether to include footer content in extraction
            extract_tables: Whether to include table content in extraction
        """
        super().__init__(sandbox)
        self.extract_headers = extract_headers
        self.extract_footers = extract_footers
        self.extract_tables = extract_tables
    
    def _is_encrypted(self, file_path: Path) -> bool:
        """
        Check if the .docx file is encrypted/password-protected.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            True if the file appears to be encrypted, False otherwise
        """
        try:
            # Try to open the document - python-docx will fail if encrypted
            Document(str(file_path))
            return False
        except Exception as e:
            # Check for encryption-related error messages
            error_msg = str(e).lower()
            if 'encrypted' in error_msg or 'password' in error_msg:
                return True
            return False
    
    def _extract_paragraphs(self, doc: Document) -> str:
        """
        Extract text from all paragraphs in the document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            String containing all paragraph text separated by newlines
        """
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        
        return '\n'.join(paragraphs)
    
    def _extract_tables(self, doc: Document) -> str:
        """
        Extract text from all tables in the document.
        
        Tables are formatted as:
        - Each row on a new line
        - Cells within a row separated by tabs
        - Tables separated by double newlines
        
        Args:
            doc: python-docx Document object
            
        Returns:
            String containing formatted table text
        """
        if not self.extract_tables:
            return ""
        
        table_contents = []
        
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    # Extract text from cell, including nested paragraphs
                    cell_text = cell.text.strip()
                    if cell_text:
                        cells.append(cell_text)
                
                if cells:
                    # Join cells with tabs for column separation
                    rows.append('\t'.join(cells))
            
            if rows:
                # Join rows with newlines, separate tables with double newline
                table_contents.append('\n'.join(rows))
        
        return '\n\n'.join(table_contents)
    
    def _extract_headers(self, doc: Document) -> str:
        """
        Extract text from all headers in the document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            String containing all header text
        """
        if not self.extract_headers:
            return ""
        
        headers = []
        
        # Extract from section headers
        for section in doc.sections:
            # Primary header
            if section.header:
                header_text = section.header.text.strip()
                if header_text:
                    headers.append(f"[HEADER: {header_text}]")
            
            # Even page header (if exists)
            if hasattr(section, 'even_page_header') and section.even_page_header:
                even_header_text = section.even_page_header.text.strip()
                if even_header_text:
                    headers.append(f"[EVEN HEADER: {even_header_text}]")
            
            # First page header (if exists)
            if hasattr(section, 'first_page_header') and section.first_page_header:
                first_header_text = section.first_page_header.text.strip()
                if first_header_text:
                    headers.append(f"[FIRST PAGE HEADER: {first_header_text}]")
        
        return '\n'.join(headers)
    
    def _extract_footers(self, doc: Document) -> str:
        """
        Extract text from all footers in the document.
        
        Args:
            doc: python-docx Document object
            
        Returns:
            String containing all footer text
        """
        if not self.extract_footers:
            return ""
        
        footers = []
        
        # Extract from section footers
        for section in doc.sections:
            # Primary footer
            if section.footer:
                footer_text = section.footer.text.strip()
                if footer_text:
                    footers.append(f"[FOOTER: {footer_text}]")
            
            # Even page footer (if exists)
            if hasattr(section, 'even_page_footer') and section.even_page_footer:
                even_footer_text = section.even_page_footer.text.strip()
                if even_footer_text:
                    footers.append(f"[EVEN FOOTER: {even_footer_text}]")
            
            # First page footer (if exists)
            if hasattr(section, 'first_page_footer') and section.first_page_footer:
                first_footer_text = section.first_page_footer.text.strip()
                if first_footer_text:
                    footers.append(f"[FIRST PAGE FOOTER: {first_footer_text}]")
        
        return '\n'.join(footers)
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract plain text from a .docx file.
        
        Extracts text from paragraphs, tables, headers, and footers based on
        configuration options. Tables are formatted for readability with
        tabs separating columns and newlines separating rows.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Extracted plain text content with structure preserved
            
        Raises:
            ParserError: If parsing fails
            EncryptedFileError: If the file is encrypted/password-protected
            FileNotFoundError: If file does not exist
            PermissionError: If file cannot be read
        """
        # Validate input
        self.validate_input(file_path)
        
        # Check for .docx extension
        if file_path.suffix.lower() != '.docx':
            self.logger.warning(f"File {file_path} does not have .docx extension")
        
        # Check if file is encrypted
        if self._is_encrypted(file_path):
            raise EncryptedFileError(
                f"File {file_path} appears to be encrypted or password-protected. "
                "This parser does not support encrypted .docx files."
            )
        
        self.logger.info(f"Extracting text from {file_path}")
        
        try:
            # Load the document
            doc = Document(str(file_path))
            
            # Collect all text parts
            text_parts = []
            
            # Extract headers (if enabled)
            headers = self._extract_headers(doc)
            if headers:
                text_parts.append(headers)
            
            # Extract main content - paragraphs
            paragraphs = self._extract_paragraphs(doc)
            if paragraphs:
                text_parts.append(paragraphs)
            
            # Extract tables (if enabled)
            tables = self._extract_tables(doc)
            if tables:
                if text_parts:
                    text_parts.append("")  # Add separator
                text_parts.append("[TABLES]")
                text_parts.append(tables)
            
            # Extract footers (if enabled)
            footers = self._extract_footers(doc)
            if footers:
                text_parts.append(footers)
            
            # Combine all parts with appropriate spacing
            full_text = '\n\n'.join(text_parts).strip()
            
            # Sanitize output
            sanitized = self.sanitize_output(full_text)
            
            self.logger.debug(f"Extracted {len(sanitized)} characters from {file_path}")
            return sanitized
        
        except PackageNotFoundError as e:
            raise ParserError(f"Invalid or corrupted .docx file: {e}")
        except Exception as e:
            raise ParserError(f"Failed to parse .docx file {file_path}: {e}")
    
    def get_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from a .docx file.
        
        Reads core properties including author, title, subject, creation date,
        modification date, and other document statistics.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Dictionary containing metadata key-value pairs
            
        Raises:
            ParserError: If metadata extraction fails
            FileNotFoundError: If file does not exist
        """
        # Validate input
        self.validate_input(file_path)
        
        self.logger.info(f"Extracting metadata from {file_path}")
        
        try:
            # Load the document
            doc = Document(str(file_path))
            
            metadata = {}
            
            # Extract core properties
            core_props = doc.core_properties
            
            # Map core properties to metadata fields
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.category:
                metadata['category'] = core_props.category
            if core_props.comments:
                metadata['comments'] = core_props.comments
            if core_props.content_status:
                metadata['content_status'] = core_props.content_status
            if core_props.created:
                metadata['creation_date'] = core_props.created
            if core_props.identifier:
                metadata['identifier'] = core_props.identifier
            if core_props.keywords:
                metadata['keywords'] = core_props.keywords
            if core_props.language:
                metadata['language'] = core_props.language
            if core_props.last_modified_by:
                metadata['last_modified_by'] = core_props.last_modified_by
            if core_props.last_printed:
                metadata['last_printed'] = core_props.last_printed
            if core_props.modified:
                metadata['modification_date'] = core_props.modified
            if core_props.revision:
                metadata['revision'] = core_props.revision
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.version:
                metadata['version'] = core_props.version
            
            # Add document statistics
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            metadata['section_count'] = len(doc.sections)
            
            # Calculate approximate character count from all paragraphs
            char_count = sum(len(p.text) for p in doc.paragraphs)
            metadata['character_count'] = char_count
            
            # Add file information
            metadata['file_size_bytes'] = file_path.stat().st_size
            metadata['file_extension'] = file_path.suffix.lower()
            metadata['parser'] = 'python-docx'
            
            self.logger.debug(f"Extracted {len(metadata)} metadata fields")
            return metadata
        
        except PackageNotFoundError as e:
            raise ParserError(f"Invalid or corrupted .docx file: {e}")
        except Exception as e:
            raise ParserError(f"Failed to extract metadata from {file_path}: {e}")
    
    def supports_encryption(self) -> bool:
        """
        Check if parser supports encrypted files.
        
        Override from BaseParser. .docx files may be encrypted with password
        protection. python-docx cannot read encrypted files.
        
        Returns:
            False (encrypted .docx files are not supported)
        """
        return False
